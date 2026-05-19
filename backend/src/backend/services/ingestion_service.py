import csv
import io
import logging
import os
import uuid
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from src.backend.config.settings import get_settings
from src.backend.vectorstore.pg_vector import add_documents

logger = logging.getLogger(__name__)
settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".docx"}

# Sections that don't contain useful knowledge — stripped before chunking
_PDF_CUTOFF_MARKERS = [
    "\nReferences\n",
    "\nFurther reading\n",
    "\nExternal links\n",
    "\nSee also\n",
    "\nNotes\n",
]


def _ensure_upload_dir() -> Path:
    p = Path(settings.upload_dir)
    p.mkdir(parents=True, exist_ok=True)
    return p


def _clean_pdf_text(text: str) -> str:
    """
    Strip bibliography, references, and other non-content sections from
    extracted PDF text so they don't pollute the vector index.
    """
    for marker in _PDF_CUTOFF_MARKERS:
        idx = text.find(marker)
        if idx != -1:
            text = text[:idx]
            logger.debug("PDF text truncated at marker '%s'", marker.strip())
            break
    return text.strip()


def _extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()

    if ext in (".txt", ".md"):
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except Exception as exc:
            raise ValueError(f"Failed to process PDF: {exc}") from exc

    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            raise ValueError(f"Failed to process DOCX: {exc}") from exc

    if ext == ".csv":
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            reader = csv.DictReader(io.StringIO(text))
            rows = [", ".join(f"{k}: {v}" for k, v in row.items()) for row in reader]
            return "\n".join(rows)
        except Exception as exc:
            raise ValueError(f"Failed to process CSV: {exc}") from exc

    raise ValueError(f"Unsupported extension '{ext}'. Allowed: {', '.join(SUPPORTED_EXTENSIONS)}")


def _split_text(text: str, source: str) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_text(text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": source, "chunk_index": i},
        )
        for i, chunk in enumerate(chunks)
    ]


def ingest_file(file_bytes: bytes, original_filename: str, saved_filename: str) -> int:
    """
    Runs full parsing, chunking, and vector insertion pipeline.
    Returns total processed chunk count.
    """
    ext = Path(original_filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported format '{ext}'. Allowed: {', '.join(SUPPORTED_EXTENSIONS)}")

    upload_path = _ensure_upload_dir() / saved_filename
    upload_path.write_bytes(file_bytes)
    logger.info("File saved to disk: %s", upload_path)

    raw_text = _extract_text(file_bytes, original_filename)
    if not raw_text.strip():
        raise ValueError("File contains no extractable text.")

    if ext == ".pdf":
        original_length = len(raw_text)
        raw_text = _clean_pdf_text(raw_text)
        logger.info(
            "PDF cleaned: %d → %d chars (removed %d chars of references/bibliography)",
            original_length,
            len(raw_text),
            original_length - len(raw_text),
        )

    docs = _split_text(raw_text, source=saved_filename)
    count = add_documents(docs)
    logger.info("Ingestion complete: %d chunks generated from '%s'", count, original_filename)
    return count


def delete_file_from_disk(saved_filename: str) -> None:
    path = Path(settings.upload_dir) / saved_filename
    if path.exists():
        os.remove(path)
        logger.info("Deleted file from disk: %s", saved_filename)


def generate_saved_filename(original_filename: str) -> str:
    ext = Path(original_filename).suffix.lower()
    return f"{uuid.uuid4().hex}{ext}"